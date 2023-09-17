import json
import tarfile
import urllib
from collections import OrderedDict
from pathlib import Path
from docx import Document
import sys
import yaml
import markdown
import os
import datetime
import re
import unidecode
import shutil
import urllib.error
import urllib.request

config = {
    "directories": {
        "input": "source",
        "output": "out",
        "themes": "themes",
    },
    "site": {
        "theme": "default",
        "domain": "wordflow.com"
    },
    "author": {
        "nickname": "ahr",
        "name": "author",
        "email": "you@me.com",
        "about": "I publish my word documents using wordflow!",
    },
    "generator": {
        "mode": "md",  # md or docx
    }
}

content = {}

theme = {}


def slugify(text):
    text = unidecode.unidecode(text).lower()
    r = re.sub(r'[\W_]+', '-', text)
    if r.endswith("-"):
        r = r[:len(r) - 1]
    return r


def html2text(htm):
    regex = re.compile(r'<[^>]+>')
    return regex.sub(' ', htm)


def docx2html(docx_file):
    doc = Document(docx_file)
    html = ""
    html += "<style>"
    for style in doc.styles:
        if style.name != "Normal":
            font_style = ""
            try:
                font_style += style.font.bold and "font-weight: bold;" or ""
                font_style += style.font.italic and "font-style: italic;" or ""
                font_size = style.font.size
                if font_size:
                    font_style += f"font-size: {font_size.pt}pt;"
            except AttributeError:
                pass  # Handle styles without font attributes

            if font_style:
                html += f".{style.name} {{ {font_style} }}\n"
    html += "</style>"

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if style_name != "Normal":
            html += f'<p class="{style_name}">{paragraph.text}</p>'
        else:
            html += f"<p>{paragraph.text}</p>"

    for table in doc.tables:
        html += "<table>"
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                html += f"<td>{cell.text}</td>"
            html += "</tr>"
        html += "</table>"

    for idx, image in enumerate(doc.inline_shapes):
        image_data = image.blob
        image_format = image.ext
        image_file = f"image_{idx}.{image_format}"
        with open(image_file, "wb") as img_file:
            img_file.write(image_data)
        html += f'<img src="{image_file}" alt="Image" />'
    return html


def md2html(markdown_file):
    with open(markdown_file, "r", encoding="utf-8") as md_file:
        markdown_content = md_file.read()

    html_content = markdown.markdown(markdown_content)

    return html_content


def scandir(directory_path):
    global content
    global config

    if not os.path.exists(directory_path):
        raise ValueError(f'Directory not found: {directory_path}')

    for filename in os.listdir(directory_path):
        if config["generator"]["mode"] == "md":
            if filename.endswith('.md'):
                markdown_file = os.path.join(directory_path, filename)
                html_content = md2html(markdown_file)
                content[filename] = {}
                content[filename]["body"] = html_content
                content[filename]["title"] = os.path.splitext(filename)[0]
                content[filename]["outfile"] = config["directories"]["output"] + "/post/" + os.path.splitext(filename)[
                    0] + ".html"
                content[filename]["date"] = datetime.datetime.fromtimestamp(os.path.getctime(markdown_file)).strftime(
                    '%Y-%m-%d %H:%M:%S')
        elif config["generator"]["mode"] == "docx":
            if filename.endswith('.docx'):
                docxfile = os.path.join(directory_path, filename)
                html_content = md2html(docxfile)
                content[filename] = {}
                content[filename]["body"] = html_content
                content[filename]["title"] = os.path.splitext(filename)[0]
                content[filename]["outfile"] = config["directories"]["output"] + "/post/" + os.path.splitext(filename)[
                    0] + ".html"
                content[filename]["date"] = datetime.datetime.fromtimestamp(os.path.getctime(docxfile)).strftime(
                    '%Y-%m-%d %H:%M:%S')
        else:
            print("mode not supported")
        with open('api.json', 'w') as f:
            f.write(json.dumps(content, indent=4))


def loadtheme():
    global theme, config
    if os.path.exists(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/config.yaml"):
        path = config["directories"]["themes"] + "/" + config["site"]["theme"]
        with open(path + "/config.yaml") as file:
            try:
                theme = yaml.safe_load(file)
            except yaml.YAMLError as exception:
                print(exception)
    else:
        print("Theme configuration file not found!")


def parsetemplate(input_data, template_type):
    global config
    loadtheme()
    theme_directory = config['directories']['themes']
    theme_name = config['site']['theme']
    theme_file = os.path.join(theme_directory, theme_name, f"{template_type}.html")

    if os.path.exists(theme_file):
        with open(theme_file, 'r', encoding='utf-8') as file:
            template_content = file.read()
        placeholder_pattern = r'{{(.*?)}}'
        matches = re.findall(placeholder_pattern, template_content)
        for match in matches:
            placeholder = f'{{{{{match}}}}}'
            if match in input_data:
                template_content = template_content.replace(placeholder, str(input_data[match]))
            else:
                print(f"Warning: Placeholder '{match}' not found in input data.")

        return template_content
    else:
        print("Warning: Template not found.")
    return None


def parsesnippet(input, snippet):
    global theme
    p = re.compile(r'\{\{([a-z]+)\}\}')
    matches = p.findall(theme["snippets"][snippet])
    output = theme["snippets"][snippet]

    for match in matches:
        if match in input:
            output = output.replace(f'{{{{{match}}}}}', str(input[match]))

    return output


def gencontent():
    global content, config
    loadtheme()
    if os.path.exists(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/assets"):
        print("Found theme assets, Copying them.")
        if os.path.exists(config["directories"]["output"] + "/public/assets"):
            shutil.rmtree(config["directories"]["output"] + "/public/assets")
        shutil.copytree(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/assets",
                        config["directories"]["output"] + "/public/assets")
    if not os.path.exists(config["directories"]["output"]):
        os.mkdir(config["directories"]["output"])
    if not os.path.exists(config["directories"]["output"] + "/post"):
        os.mkdir(config["directories"]["output"] + "/post")
    for document in content:
        document = content[document]
        document.update(config["author"])
        document.update(config["site"])
        outfile = open(document["outfile"], "w")
        outfile.write(parsetemplate(document, "post"))
        outfile.close()
    if os.path.exists("api.json"):
        shutil.copy("api.json", config["directories"]["output"] + "/api.json")


def generatehomepage():
    global config
    global content
    homecontent = {}
    homecontent.update(config["author"])
    homecontent.update(config["site"])
    homecontent["body"] = ""
    tempcontent = {}
    date_order = OrderedDict(sorted(content.items()), key=lambda t: t["date"])
    for post in date_order:
        if type(date_order[post]) is dict:
            tempcontent.update(content[post])
            tempcontent["file"] = str.replace(tempcontent["outfile"], config["directories"]["output"] + "/", "")
            tempcontent["body"] = html2text(tempcontent["body"])
            tempcontent["body"] = (tempcontent["body"][:120] + '..') if len(tempcontent["body"]) > 120 else \
                tempcontent["body"]
            homecontent["body"] += parsesnippet(tempcontent, "home_post")
            tempcontent = {}
        filename = config["directories"]["output"] + "/index.html"
        outfile = open(filename, "w")
        outfile.write(parsetemplate(homecontent, "home"))
        outfile.close()


def downloadtheme(name):
    global config
    url = "https://api.github.com/repos/devsimsek/WordFlow_themes/tarball/" + name + "_theme"
    try:
        status = urllib.request.urlopen(url)
    except urllib.error.HTTPError:
        print("Theme " + name + " not found.")
        return
    if not os.path.exists(config["directories"]["themes"] + "/" + name):
        if not os.path.exists("temp"):
            os.mkdir("temp")
        if not os.path.exists(config["directories"]["themes"] + "/" + name):
            os.mkdir(config["directories"]["themes"] + "/" + name)
        urllib.request.urlretrieve(url, "temp/" + name + "_theme.tar.gz")
        theme = tarfile.open("temp/" + name + "_theme.tar.gz")
        theme.extractall(config["directories"]["themes"] + "/" + name)
        extractedfile = os.path.commonprefix(theme.getnames())
        theme.close()
        for file in Path(config["directories"]["themes"] + "/" + name + "/" + extractedfile).glob("*"):
            shutil.move(file, config["directories"]["themes"] + "/" + name)
        shutil.rmtree(config["directories"]["themes"] + "/" + name + "/" + extractedfile)
        shutil.rmtree("temp")
        print("Theme installed.")
    else:
        print("Selected theme already exists. Want to reinstall?")
        val = input("(yes, no)> ")
        if val != "yes":
            print("Skipping theme installation.")
        else:
            print("Reinstalling...")
            shutil.rmtree(config["directories"]["themes"] + "/" + name)
            downloadtheme(name)


def clearinstallation():
    global config
    for directory in config["directories"]:
        print("removing " + config["directories"][directory])
        shutil.rmtree(config["directories"][directory])
    if os.path.exists("config.yaml"):
        os.remove("config.yaml")
    if os.path.exists("generated_output.json"):
        os.remove("generated_output.json")


def clearcontent():
    global config
    for directory in config["directories"]:
        if directory == "input":
            continue
        if directory == "themes":
            continue
        print("removing " + config["directories"][directory])
        shutil.rmtree(config["directories"][directory])
        if not os.path.exists(config["directories"][directory]):
            os.mkdir(config["directories"][directory])
        if os.path.exists("api.json"):
            os.remove("api.json")


def initapp():
    """
    Initialize wordflow application
    """
    global config
    print("Welcome to the WordFlow initializer.")
    print("Checking configuration")
    if not os.path.exists("config.yaml"):
        for key in config:
            i = 1
            for opt in config[key]:
                print(
                    "--- " + key.capitalize() + " Configuration (" + str(i) + " of " + str(len(config[key])) + ") ---")
                print("Configuring " + opt + " field.")
                val = input("Value (default: " + config[key][opt] + ")> ")
                if not val == "":
                    config[key][opt] = val
                i += 1
        print("Configuration completed.")
        if not os.path.exists("config.yaml"):
            with open("config.yaml", "w") as file:
                try:
                    yaml.dump(config, file)
                    print("Configuration saved. You can create your documents now.")
                except yaml.YAMLError as exception:
                    print(exception)
        else:
            print("Operation failed. Configuration already exists!")
            val = input(
                "Want to clean install WordFlow? (this will remove every configuration and files.) (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearinstallation()
    else:
        print("Configuration found. Skipping.")
    print("Checking directories")
    for directory in config["directories"]:
        if not os.path.exists(config["directories"][directory]):
            print(config["directories"][directory] + " not exists. Creating.")
            os.mkdir(config["directories"][directory])
        else:
            print(directory + " exists.")
    if os.path.exists(config["directories"]["input"]):
        if not os.path.exists(config["directories"]["input"] + "/post"):
            os.mkdir(config["directories"]["input"] + "/post")
        if not os.path.exists(config["directories"]["input"] + "/page"):
            os.mkdir(config["directories"]["input"] + "/page")
    print("Installing Theme")
    downloadtheme(config["site"]["theme"])

    print("Application should be initialized correctly. Thanks for using WordFlow.")
    exit(0)


def wordflow():
    if not os.path.exists("config.yaml"):
        print("Warning: Configuration file not found. Launching initializer.")
        initapp()
    else:
        global config
        with open("config.yaml") as file:
            try:
                config = yaml.safe_load(file)
            except yaml.YAMLError as exception:
                print(exception)


def argvparser():
    global config
    args = sys.argv[1:]
    force = False
    for arg in args:
        if arg == "gen" or arg == "-g":
            print("Started scan.")
            scandir(config["directories"]["input"])
            gencontent()
            generatehomepage()
        elif arg == "-f" or arg == "--force":
            force = True
        elif arg == "init" or arg == "-init":
            initapp()
        elif arg == "clear" or arg == "-c":
            val = input(
                "Want to clean install WordFlow? (this will remove every configuration and files.) (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearinstallation()
        elif arg == "clear":
            val = input(
                "Want to clean install WordFlow? (this will remove every configuration and files.) (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearinstallation()
        elif arg == "installtheme" or arg == "theme":
            name = input("Theme name> ")
            downloadtheme(name)
        elif arg == "clearcontent" or arg == "-cc":
            if not force:
                val = input(
                    "Want to wipe all generated content? (yes or no)> ")
                if val != "yes":
                    print("Bye :)")
                else:
                    clearcontent()
            else:
                clearcontent()


if __name__ == "__main__":
    wordflow()
    argvparser()
else:
    print("illegal launch option")
